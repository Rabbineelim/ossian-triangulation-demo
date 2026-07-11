"""Format-preserving cleaned-file export.

Whatever format a source was uploaded in, the reviewer can download its cleaned
version *individually in the same format*:

    csv/tsv/xlsx  -> rebuilt table (removed rows dropped, content column cleaned,
                     dates standardized), same delimiter / workbook format
    txt           -> cleaned text, section headers preserved
    vtt           -> cleaned WebVTT cues (timestamps kept)
    docx          -> cleaned paragraphs in a Word document
    pdf           -> cleaned paragraphs in a PDF

Only kept units are written (empty + duplicate units are the ones the rule engine
removed). The original upload always remains untouched in storage/ (evidence trail).
"""
from __future__ import annotations

import csv
import io
from pathlib import Path
from typing import Any

import pandas as pd

from . import db
from .clean import rules
from .ingest.extractors import _DATE_HINT, _pick_text_column

MEDIA = {
    "csv": "text/csv",
    "tsv": "text/tab-separated-values",
    "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    "txt": "text/plain; charset=utf-8",
    "vtt": "text/vtt; charset=utf-8",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "doc": "application/msword",
    "pdf": "application/pdf",
}


def source_is_downloadable(file_type: str) -> bool:
    return file_type in MEDIA


def _content_column(conn, source_id: int, df: pd.DataFrame) -> str | None:
    """The column that holds the free text: prefer the mapping saved at import,
    else re-derive it deterministically the same way extraction did."""
    rows = conn.execute(
        "SELECT suggestion_text FROM ai_suggestions WHERE source_id=? "
        "AND suggestion_type='column_mapping'", (source_id,)).fetchall()
    for r in rows:
        txt = r["suggestion_text"] or ""
        if "->" in txt:
            inc, std = (p.strip() for p in txt.split("->", 1))
            if std == "content" and inc in df.columns:
                return inc
    return _pick_text_column(df) if len(df.columns) else None


def _load_kept(conn, source_id: int) -> list[dict[str, Any]]:
    rows = conn.execute(
        """SELECT ru.unit_id, cu.cleaned_text, ru.metadata, ru.row_number,
                  ru.page_number, cu.quality_flags
           FROM cleaned_units cu JOIN raw_units ru ON ru.unit_id = cu.unit_id
           WHERE ru.source_id=? AND cu.status='kept'
           ORDER BY ru.unit_id""", (source_id,)).fetchall()
    return [
        {"unit_id": r["unit_id"], "cleaned_text": r["cleaned_text"] or "",
         "metadata": db.loads(r["metadata"]) or {},
         "row_number": r["row_number"], "page_number": r["page_number"],
         "flags": db.loads(r["quality_flags"]) or []}
        for r in rows
    ]


# --- tabular ---------------------------------------------------------------
def _build_table(conn, source_id: int, kept: list[dict]) -> pd.DataFrame:
    rows, cleaned = [], []
    for u in kept:
        md = {k: v for k, v in u["metadata"].items() if k != "index"}
        rows.append(md)
        cleaned.append(u["cleaned_text"])
    df = pd.DataFrame(rows)
    if df.empty:
        return df
    text_col = _content_column(conn, source_id, df)
    if text_col and text_col in df.columns:
        df[text_col] = cleaned
    # standardize date-like columns to one consistent (ISO) format
    for col in df.columns:
        if _DATE_HINT.search(str(col)):
            df[col] = df[col].map(lambda v: (rules.standardize_date(v)[0] or v)
                                  if v not in (None, "") else v)
    return df


def _tabular_bytes(df: pd.DataFrame, file_type: str) -> bytes:
    if file_type == "xlsx":
        buf = io.BytesIO()
        with pd.ExcelWriter(buf, engine="openpyxl") as xl:
            df.to_excel(xl, index=False, sheet_name="cleaned")
        return buf.getvalue()
    sep = "\t" if file_type == "tsv" else ","
    text = df.to_csv(index=False, sep=sep, quoting=csv.QUOTE_MINIMAL)
    return text.encode("utf-8-sig")  # BOM so Excel opens UTF-8 correctly


# --- text formats ----------------------------------------------------------
def _txt_bytes(kept: list[dict]) -> bytes:
    out, last_section = [], None
    for u in kept:
        section = u["metadata"].get("section")
        if section and section != last_section:
            out.append(f"\n===== {section} =====\n")
            last_section = section
        marker = u["metadata"].get("marker")
        if marker:
            out.append(f"[{marker}]")
        out.append(u["cleaned_text"])
        out.append("")  # blank line between units
    return ("\n".join(out).strip() + "\n").encode("utf-8")


def _vtt_bytes(kept: list[dict]) -> bytes:
    lines = ["WEBVTT", ""]
    n = 0
    for u in kept:
        text = u["cleaned_text"].strip()
        if not text:
            continue
        n += 1
        lines.append(str(n))
        start, end = u["metadata"].get("start"), u["metadata"].get("end")
        if start and end:
            lines.append(f"{start} --> {end}")
        lines.append(text)
        lines.append("")
    return ("\n".join(lines)).encode("utf-8")


def _docx_bytes(kept: list[dict], title: str) -> bytes:
    import docx
    doc = docx.Document()
    doc.add_heading(f"Cleaned — {title}", level=1)
    last_section = None
    for u in kept:
        section = u["metadata"].get("section")
        if section and section != last_section:
            doc.add_heading(section, level=2)
            last_section = section
        text = u["cleaned_text"].strip()
        if text:
            doc.add_paragraph(text)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _doc_paragraphs(kept: list[dict], title: str) -> list[str]:
    paras = [f"Cleaned — {title}"]
    last_section = None
    for u in kept:
        section = u["metadata"].get("section")
        if section and section != last_section:
            paras.append(section)
            last_section = section
        text = u["cleaned_text"].strip()
        if text:
            paras.append(text)
    return paras


def _doc_bytes(kept: list[dict], title: str) -> tuple[bytes, str]:
    """Write a true .doc via Word; fall back to .docx. Returns (data, ext)."""
    import os
    import tempfile

    from .ingest.legacy_doc import can_write_doc, write_doc
    paras = _doc_paragraphs(kept, title)
    if can_write_doc():
        fd, tmp = tempfile.mkstemp(suffix=".doc")
        os.close(fd)
        try:
            if write_doc(paras, tmp) and os.path.getsize(tmp) > 0:
                data = Path(tmp).read_bytes()
                return data, "doc"
        finally:
            try:
                os.unlink(tmp)
            except OSError:
                pass
    # fallback: deliver a .docx (format changed, but content is identical)
    return _docx_bytes(kept, title), "docx"


def _pdf_bytes(kept: list[dict], title: str) -> bytes:
    from fpdf import FPDF

    def latin(s: str) -> str:  # core fonts are latin-1; sanitize gracefully
        return (s or "").encode("latin-1", "replace").decode("latin-1")

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 14)
    pdf.multi_cell(0, 8, latin(f"Cleaned — {title}"))
    pdf.ln(2)
    pdf.set_font("Helvetica", size=11)
    last_section = None
    for u in kept:
        section = u["metadata"].get("section")
        if section and section != last_section:
            pdf.set_font("Helvetica", "B", 12)
            pdf.multi_cell(0, 7, latin(section))
            pdf.set_font("Helvetica", size=11)
            last_section = section
        text = u["cleaned_text"].strip()
        if text:
            pdf.multi_cell(0, 6, latin(text))
            pdf.ln(1)
    out = pdf.output(dest="S")
    return bytes(out) if isinstance(out, (bytes, bytearray)) else out.encode("latin-1")


def build_project_zip(conn, project_id: int) -> bytes:
    """Bundle every cleaned source, each in its OWN original format, into a ZIP."""
    import zipfile
    srcs = conn.execute(
        "SELECT source_id, file_name, file_type, import_status, n_units "
        "FROM sources WHERE project_id=? ORDER BY source_id", (project_id,)).fetchall()
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as z:
        added = 0
        for s in srcs:
            if s["import_status"] != "cleaned" or not source_is_downloadable(s["file_type"]):
                continue
            try:
                data, fname, _ = build_cleaned_file(conn, s["source_id"])
            except Exception:
                continue
            z.writestr(f"{s['source_id']:03d}_{fname}", data)
            added += 1
        if added == 0:
            z.writestr("README.txt",
                       b"No cleaned sources yet. Run cleaning first, then export.")
    return buf.getvalue()


# --- entry point -----------------------------------------------------------
def build_cleaned_file(conn, source_id: int) -> tuple[bytes, str, str]:
    """Return (data, filename, media_type) for the cleaned source in its own format."""
    src = conn.execute("SELECT * FROM sources WHERE source_id=?", (source_id,)).fetchone()
    if src is None:
        raise ValueError(f"source {source_id} not found")
    file_type = src["file_type"]
    stem = src["file_name"].rsplit(".", 1)[0]
    kept = _load_kept(conn, source_id)

    if file_type in ("csv", "tsv", "xlsx"):
        df = _build_table(conn, source_id, kept)
        data = _tabular_bytes(df, file_type)
    elif file_type == "vtt":
        data = _vtt_bytes(kept)
    elif file_type == "docx":
        data = _docx_bytes(kept, stem)
    elif file_type == "doc":
        data, file_type = _doc_bytes(kept, stem)  # may fall back to docx
    elif file_type == "pdf":
        data = _pdf_bytes(kept, stem)
    else:  # txt and any text-like fallback
        data = _txt_bytes(kept)
        file_type = file_type if file_type in MEDIA else "txt"

    ext = file_type if file_type in MEDIA else "txt"
    filename = f"cleaned_{stem}.{ext}"
    return data, filename, MEDIA.get(ext, "application/octet-stream")
