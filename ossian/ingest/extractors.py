"""Extraction module (Step 2, question 3: "How can the content be extracted?").

Each extractor turns one file into an ExtractionResult carrying analyzable
units. Guide rules honoured here:
    * one review / survey row / interview paragraph / transcript cue = one unit
    * keep speaker labels when available
    * preserve the raw extracted text (cleaning happens later in Step 3)
    * every unit keeps row/page/metadata for the evidence trail
"""
from __future__ import annotations

import csv
import io
import re
from pathlib import Path
from typing import Any

import pandas as pd

from .. import config
from ..models import ExtractedUnit, ExtractionResult

# --------------------------------------------------------------------------
# encoding helpers (chardet not installed -> deterministic fallback ladder)
# --------------------------------------------------------------------------
_ENCODINGS = ["utf-8-sig", "utf-8", "cp1252", "latin-1"]


def _read_text(path: Path) -> tuple[str, str]:
    raw = path.read_bytes()
    for enc in _ENCODINGS:
        try:
            return raw.decode(enc), enc
        except UnicodeDecodeError:
            continue
    return raw.decode("latin-1", errors="replace"), "latin-1(replace)"


def _sniff_sep(sample: str, default: str) -> str:
    try:
        dialect = csv.Sniffer().sniff(sample, delimiters=",\t;|")
        return dialect.delimiter
    except csv.Error:
        return default


# --------------------------------------------------------------------------
# TEXT: interviews, chatbot logs, reddit posts, diary/wellness logs
# --------------------------------------------------------------------------
_SEP_RE = re.compile(r"^[\s=_\-*#~]{6,}$")              # ==== / ---- separator lines
_SECTION_RE = re.compile(
    r"^(PARTICIPANT|CONVERSATION|POST|SESSION|ENTRY|INTERVIEW|DAY)\b.*", re.I
)
_MARKER_RE = re.compile(r"^\[[^\]]{1,80}\]$")           # [Task: ...] / [Overall reflection]
_SPEAKER_RE = re.compile(r"^([A-Z][\w .,'&/-]{0,40}?):\s+(.*)$")  # "User: ...", "Majie ZELLER: ..."


def extract_txt(path: Path, file_type: str = "txt") -> ExtractionResult:
    text, enc = _read_text(path)
    lines = text.splitlines()

    # Split into blank-line-separated blocks, tracking the current section header.
    units: list[ExtractedUnit] = []
    section: str | None = None
    marker: str | None = None
    block: list[str] = []
    warnings: list[str] = []
    header_zone = True  # leading study-metadata lines before the first separator

    def flush(sec: str | None, mark: str | None) -> None:
        nonlocal block
        joined = "\n".join(block).strip()
        block = []
        if not joined:
            return
        meta: dict[str, Any] = {}
        if sec:
            meta["section"] = sec
        if mark:
            meta["marker"] = mark  # e.g. the task being performed
        m = _SPEAKER_RE.match(joined.replace("\n", " ").strip())
        if m and len(m.group(1)) <= 40:
            meta["speaker"] = m.group(1).strip()
        units.append(ExtractedUnit(original_text=joined, metadata=meta))

    for raw_line in lines:
        line = raw_line.rstrip()
        if _SEP_RE.match(line):
            flush(section, marker)
            continue
        if _SECTION_RE.match(line.strip()):
            flush(section, marker)
            section = line.strip()
            marker = None
            header_zone = False
            continue
        if _MARKER_RE.match(line.strip()):
            # Structural marker ([Task: ...]) -> context for following units,
            # not evidence itself.
            flush(section, marker)
            marker = line.strip().strip("[]")
            header_zone = False
            continue
        if not line.strip():
            flush(section, marker)
            continue
        if header_zone and ":" in line and len(line) < 80 and not _SPEAKER_RE.match(line):
            # Study metadata line (e.g. "Method: ...", "Date: ...") — keep as note.
            continue
        block.append(line)
    flush(section, marker)

    if not units:
        warnings.append("No text blocks extracted — file may be empty or unstructured.")

    return ExtractionResult(
        file_name=path.name, file_type=file_type, detected_by="",
        source_type="", source_type_confidence=0.0,
        units=units, warnings=warnings,
    )


# --------------------------------------------------------------------------
# VTT / SRT transcripts (from the video usability study)
# --------------------------------------------------------------------------
_TS_RE = re.compile(r"(\d{2}:\d{2}:\d{2}[.,]\d{3})\s*-->\s*(\d{2}:\d{2}:\d{2}[.,]\d{3})")


def extract_vtt(path: Path, file_type: str = "vtt") -> ExtractionResult:
    text, _ = _read_text(path)
    blocks = re.split(r"\n\s*\n", text)
    units: list[ExtractedUnit] = []
    for blk in blocks:
        rows = [r for r in blk.splitlines() if r.strip()]
        if not rows or rows[0].strip().upper() == "WEBVTT":
            continue
        start = end = None
        text_lines: list[str] = []
        for r in rows:
            m = _TS_RE.search(r)
            if m:
                start, end = m.group(1), m.group(2)
            elif r.strip().isdigit() and start is None:
                continue  # cue index
            else:
                text_lines.append(r.strip())
        content = " ".join(text_lines).strip()
        if not content:
            continue
        meta: dict[str, Any] = {}
        if start:
            meta["start"], meta["end"] = start, end
        sm = _SPEAKER_RE.match(content)
        if sm:
            meta["speaker"] = sm.group(1).strip()
        units.append(ExtractedUnit(original_text=content, metadata=meta))
    return ExtractionResult(
        file_name=path.name, file_type=file_type, detected_by="",
        source_type="", source_type_confidence=0.0, units=units,
    )


# --------------------------------------------------------------------------
# PDF
# --------------------------------------------------------------------------
def extract_pdf(path: Path, file_type: str = "pdf") -> ExtractionResult:
    units: list[ExtractedUnit] = []
    warnings: list[str] = []
    try:
        import pdfplumber
        with pdfplumber.open(str(path)) as pdf:
            for pno, page in enumerate(pdf.pages, start=1):
                page_text = page.extract_text() or ""
                for para in re.split(r"\n\s*\n", page_text):
                    para = para.strip()
                    if para:
                        units.append(ExtractedUnit(
                            original_text=para, page_number=pno,
                            metadata={"page": pno},
                        ))
    except Exception as exc:  # pragma: no cover - depends on file
        warnings.append(f"PDF extraction error: {exc}")
    return ExtractionResult(
        file_name=path.name, file_type=file_type, detected_by="",
        source_type="", source_type_confidence=0.0,
        units=units, warnings=warnings,
    )


# --------------------------------------------------------------------------
# DOCX
# --------------------------------------------------------------------------
def extract_docx(path: Path, file_type: str = "docx") -> ExtractionResult:
    units: list[ExtractedUnit] = []
    warnings: list[str] = []
    try:
        import docx
        doc = docx.Document(str(path))
        for para in doc.paragraphs:
            t = para.text.strip()
            if t:
                meta: dict[str, Any] = {}
                sm = _SPEAKER_RE.match(t)
                if sm:
                    meta["speaker"] = sm.group(1).strip()
                units.append(ExtractedUnit(original_text=t, metadata=meta))
        for ti, tbl in enumerate(doc.tables):
            for ri, row in enumerate(tbl.rows):
                cells = [c.text.strip() for c in row.cells]
                joined = " | ".join(c for c in cells if c)
                if joined:
                    units.append(ExtractedUnit(
                        original_text=joined, row_number=ri,
                        metadata={"table": ti, "row": ri},
                    ))
    except Exception as exc:  # pragma: no cover
        warnings.append(f"DOCX extraction error: {exc}")
    return ExtractionResult(
        file_name=path.name, file_type=file_type, detected_by="",
        source_type="", source_type_confidence=0.0,
        units=units, warnings=warnings,
    )


# --------------------------------------------------------------------------
# Legacy DOC (Word 97-2003 / RTF)
# --------------------------------------------------------------------------
def extract_doc(path: Path, file_type: str = "doc") -> ExtractionResult:
    from .legacy_doc import read_doc_text
    warnings: list[str] = []
    text, method = read_doc_text(path)
    units: list[ExtractedUnit] = []
    if not text:
        warnings.append(
            "Could not read this legacy .doc. Install Microsoft Word (best) or "
            "re-save the file as .docx. The original is stored for the evidence trail.")
        return ExtractionResult(
            file_name=path.name, file_type=file_type, detected_by="",
            source_type="Document", source_type_confidence=0.4,
            units=units, warnings=warnings, extractable=False)

    warnings.append(f"Read .doc via {method}.")
    # Word uses \r for paragraph breaks and \x07 for table cell/row ends.
    for para in re.split(r"[\r\n\x07\x0b\x0c]+", text):
        para = para.strip()
        if not para:
            continue
        meta: dict[str, Any] = {}
        sm = _SPEAKER_RE.match(para)
        if sm and len(sm.group(1)) <= 40:
            meta["speaker"] = sm.group(1).strip()
        units.append(ExtractedUnit(original_text=para, metadata=meta))

    if not units:
        warnings.append("No text blocks extracted from the .doc.")
    return ExtractionResult(
        file_name=path.name, file_type=file_type, detected_by="",
        source_type="", source_type_confidence=0.0,
        units=units, warnings=warnings,
    )


# --------------------------------------------------------------------------
# TABULAR: CSV / TSV / XLSX
# --------------------------------------------------------------------------
_ID_HINT = re.compile(r"(^|_)(id|uid|guid|key|index)$", re.I)
_DATE_HINT = re.compile(r"(date|time|timestamp|created|updated)", re.I)


def _normalise_columns(cols: list[str]) -> tuple[list[str], bool]:
    """Rename junk headers (empty / 'Unnamed: 0' / '0','1',...) to stable names.

    Returns (new_names, junk_header_detected).
    """
    junk = 0
    out: list[str] = []
    for i, c in enumerate(cols):
        cs = str(c).strip()
        if cs == "" or cs.lower().startswith("unnamed") or cs.isdigit():
            junk += 1
            out.append("index" if i == 0 and (cs == "" or cs.lower().startswith("unnamed"))
                       else f"col_{i}")
        else:
            out.append(cs)
    return out, junk >= max(2, len(cols) // 2)


def _pick_text_column(df: pd.DataFrame) -> str | None:
    """Choose the primary free-text column: highest mean length, alphabetic, not id/date."""
    best, best_score = None, 0.0
    for col in df.columns:
        if _ID_HINT.search(str(col)) or _DATE_HINT.search(str(col)):
            continue
        series = df[col].dropna().astype(str)
        if series.empty:
            continue
        mean_len = series.str.len().mean()
        alpha_frac = series.str.contains(r"[A-Za-z]").mean()
        score = mean_len * alpha_frac
        if score > best_score and mean_len >= 12:
            best, best_score = col, score
    return best


def _load_table(path: Path, file_type: str) -> tuple[pd.DataFrame, str, list[str]]:
    warnings: list[str] = []
    if file_type == "xlsx":
        df = pd.read_excel(path, dtype=str)
        return df, "utf-8(xlsx)", warnings

    text, enc = _read_text(path)
    default_sep = "\t" if file_type == "tsv" else ","
    sep = _sniff_sep("\n".join(text.splitlines()[:10]), default_sep)

    raw_line_count = text.count("\n")
    # Robust parse: skip malformed rows but remember we did (ragged real-world files).
    df = pd.read_csv(
        io.StringIO(text), sep=sep, engine="python", dtype=str,
        on_bad_lines="skip", quoting=csv.QUOTE_MINIMAL,
        nrows=None,
    )
    loaded = len(df)
    approx_skipped = max(0, raw_line_count - 1 - loaded)
    if approx_skipped > 0:
        warnings.append(
            f"~{approx_skipped} malformed/ragged rows skipped during extraction "
            f"(raw lines={raw_line_count}, parsed rows={loaded})."
        )
    return df, enc, warnings


def extract_table(path: Path, file_type: str) -> ExtractionResult:
    df, enc, warnings = _load_table(path, file_type)

    new_cols, junk = _normalise_columns(list(df.columns))
    df.columns = new_cols
    if junk:
        warnings.append("Junk/generic header detected (blank or numeric column names) "
                        "— columns renamed; mapping suggested by AI layer.")

    truncated = False
    if len(df) > config.MAX_UNITS_PER_SOURCE:
        df = df.head(config.MAX_UNITS_PER_SOURCE)
        truncated = True

    text_col = _pick_text_column(df)
    mapping: dict[str, str] = {}
    if text_col:
        mapping[text_col] = "content"

    units: list[ExtractedUnit] = []
    for ridx, (_, row) in enumerate(df.iterrows()):
        rowd = {k: (None if pd.isna(v) else str(v)) for k, v in row.items()}
        content = (rowd.get(text_col) or "") if text_col else " | ".join(
            str(v) for v in rowd.values() if v)
        units.append(ExtractedUnit(
            original_text=content or "",
            row_number=ridx,
            metadata=rowd,
        ))

    return ExtractionResult(
        file_name=path.name, file_type=file_type, detected_by="",
        source_type="", source_type_confidence=0.0,
        units=units, columns=new_cols, column_mapping=mapping,
        warnings=warnings, truncated=truncated,
    )


# --------------------------------------------------------------------------
# Roadmap formats (audio / video / image) — accepted but not yet extracted
# --------------------------------------------------------------------------
def extract_roadmap(path: Path, file_type: str) -> ExtractionResult:
    return ExtractionResult(
        file_name=path.name, file_type=file_type, detected_by="",
        source_type="Media", source_type_confidence=1.0, units=[],
        warnings=[f"'{file_type}' is a future-roadmap format. The original is "
                  "stored for the evidence trail; extraction arrives after the "
                  "import foundation is stable (per Step 2 guide)."],
        extractable=False,
    )


# --------------------------------------------------------------------------
# dispatch
# --------------------------------------------------------------------------
_DISPATCH = {
    "txt": lambda p: extract_txt(p, "txt"),
    "vtt": lambda p: extract_vtt(p, "vtt"),
    "pdf": lambda p: extract_pdf(p, "pdf"),
    "docx": lambda p: extract_docx(p, "docx"),
    "doc": lambda p: extract_doc(p, "doc"),
    "csv": lambda p: extract_table(p, "csv"),
    "tsv": lambda p: extract_table(p, "tsv"),
    "xlsx": lambda p: extract_table(p, "xlsx"),
}


def extract(path: str | Path, file_type: str) -> ExtractionResult:
    path = Path(path)
    if file_type in _DISPATCH:
        return _DISPATCH[file_type](path)
    if file_type in ("audio", "video", "image"):
        return extract_roadmap(path, file_type)
    # Unknown / json fall back to text.
    return extract_txt(path, file_type)
